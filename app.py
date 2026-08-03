import streamlit as st
import pandas as pd
from datetime import datetime
import urllib.request
import json
import time
import unicodedata
import os
import re
import difflib
import numpy as np

# Leitores de arquivos
import docx
from pypdf import PdfReader

# Componente para gravação de áudio e Biblioteca de IA Local
from streamlit_mic_recorder import speech_to_text
from sentence_transformers import SentenceTransformer, util

st.set_page_config(page_title="Requisição Diária", page_icon="📝", layout="centered")

# ➔ SUA URL DO APP DA WEB DO GOOGLE SCRIPTS (Terminada em /exec)
URL_WEB_APP = "https://script.google.com/macros/s/AKfycbzcNped3ftP-9FkLcWC-u65kl0RlX-rW2Z_8AHLGKgrw2ETjkoKJI2CHqisiSQnoUUb/exec"

# ==============================================================================
# 1. CARREGAMENTO DO MODELO DE INTELIGÊNCIA ARTIFICIAL LOCAL
# ==============================================================================
@st.cache_resource
def carregar_modelo_ia():
    # Modelo multilíngue leve e otimizado (~120MB), ideal para o servidor do Streamlit Cloud
    return SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

modelo_ia = carregar_modelo_ia()

@st.cache_data
def calcular_embeddings_catalogo(_modelo, catalogo):
    if not catalogo:
        return None
    return _modelo.encode(catalogo, convert_to_tensor=True)

# ==============================================================================
# 2. FUNÇÕES DE PROCESSAMENTO DE TEXTO, FALA E EXTRAÇÃO DE ARQUIVOS
# ==============================================================================
def remover_acentos(texto): 
    return "".join(c for c in unicodedata.normalize('NFD', str(texto)) if unicodedata.category(c) != 'Mn').upper().strip()

def normalizar_texto(texto):
    texto_nfd = unicodedata.normalize('NFD', str(texto))
    texto_sem_acento = "".join(c for c in texto_nfd if unicodedata.category(c) != 'Mn').upper()
    texto_limpo = re.sub(r'[^A-Z0-9\s]', ' ', texto_sem_acento)
    return " ".join(texto_limpo.split())

STOP_WORDS = {
    "QUERO", "ME", "DA", "VEJA", "POR", "FAVOR", "MANDA", "COLOCA", "ADICIONA", 
    "PRECISO", "DE", "DO", "DA", "DOS", "DAS", "E", "MAIS", "TAMBEM", "GOSTARIA",
    "UNIDADE", "UNIDADES", "KILO", "KILOS", "QUILO", "QUILOS", "KG", "LATA", "LATAS", 
    "PACOTE", "PACOTES", "GARRAFA", "GARRAFAS", "CAIXA", "CAIXAS", "UND", "PCT", "CX", "GFA", "POTE"
}

NUMEROS_EXTENSO = {
    "UM": 1, "UMA": 1, "DOIS": 2, "DUAS": 2, "TRES": 3, "TRÊS": 3, "QUATRO": 4, "CINCO": 5,
    "SEIS": 6, "SETE": 7, "OITO": 8, "NOVE": 9, "DEZ": 10, "ONZE": 11, "DOZE": 12,
    "QUINZE": 15, "VINTE": 20, "TRINTA": 30, "CINQUENTA": 50
}

# Leitor Universal de Documentos (Excel, Word e PDF)
def extrair_linhas_do_arquivo(uploaded_file):
    nome = uploaded_file.name.lower()
    linhas = []
    
    if nome.endswith(('.xlsx', '.xls')):
        try:
            df = pd.read_excel(uploaded_file)
            for _, row in df.iterrows():
                linha_txt = " ".join([str(val) for val in row.values if pd.notna(val)])
                if linha_txt.strip():
                    linhas.append(linha_txt)
        except Exception as e:
            st.error(f"Erro ao ler arquivo Excel: {e}")
            
    elif nome.endswith('.docx'):
        try:
            doc = docx.Document(uploaded_file)
            for p in doc.paragraphs:
                if p.text.strip():
                    linhas.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    linha_txt = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if linha_txt.strip():
                        linhas.append(linha_txt)
        except Exception as e:
            st.error(f"Erro ao ler arquivo Word: {e}")
            
    elif nome.endswith('.pdf'):
        try:
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                texto = page.extract_text()
                if texto:
                    for l in texto.split('\n'):
                        if l.strip():
                            linhas.append(l.strip())
        except Exception as e:
            st.error(f"Erro ao ler arquivo PDF: {e}")
            
    return linhas

# Separa texto ou fala em múltiplos itens e extrai suas quantidades
def extrair_itens_da_fala(texto_falado):
    texto_norm = normalizar_texto(texto_falado)
    
    palavras = texto_norm.split()
    palavras_convertidas = []
    for p in palavras:
        if p in NUMEROS_EXTENSO:
            palavras_convertidas.append(str(NUMEROS_EXTENSO[p]))
        else:
            palavras_convertidas.append(p)
    texto_tratado = " ".join(palavras_convertidas)
    
    clausulas = re.split(r'\b(?:E|MAIS|TAMBEM)\b|[,;]', texto_tratado)
    
    sub_frases = []
    for c in clausulas:
        c = c.strip()
        if not c:
            continue
        tokens = c.split()
        corrente = []
        for token in tokens:
            if token.isdigit() and corrente:
                palavras_uteis = [w for w in corrente if w not in STOP_WORDS and not w.isdigit()]
                if palavras_uteis:
                    sub_frases.append(" ".join(corrente))
                    corrente = [token]
                else:
                    corrente.append(token)
            else:
                corrente.append(token)
        if corrente:
            sub_frases.append(" ".join(corrente))
            
    resultados = []
    for sf in sub_frases:
        tokens_sf = sf.split()
        qtd = 1
        termos = []
        for t in tokens_sf:
            if t.isdigit():
                qtd = int(t)
            elif t not in STOP_WORDS:
                termos.append(t)
        termo_busca = " ".join(termos)
        
        if termo_busca:
            resultados.append((sf, qtd, termo_busca))
            
    return resultados

# Algoritmo de Busca de Alta Precisão (IA + Substring + Matching Palavra-por-Palavra)
def encontrar_item_super_assertivo(termo_busca, catalogo_itens, embeddings_catalogo):
    if not termo_busca or not catalogo_itens or embeddings_catalogo is None:
        return None, 0.0
        
    termo_norm = normalizar_texto(termo_busca)
    words_busca = termo_norm.split()
    
    if not words_busca:
        return None, 0.0
        
    embedding_fala = modelo_ia.encode(termo_busca, convert_to_tensor=True)
    scores_ia = util.cos_sim(embedding_fala, embeddings_catalogo)[0]
    
    melhor_item = None
    melhor_score_final = 0.0
    
    for idx, item in enumerate(catalogo_itens):
        item_norm = normalizar_texto(item)
        words_item = item_norm.split()
        score_ia = float(scores_ia[idx])
        
        if termo_norm in item_norm:
            score_substring = 1.0
        else:
            score_substring = 0.0
            
        scores_palavras = []
        for pb in words_busca:
            best_p = 0.0
            for pi in words_item:
                if pb == pi:
                    best_p = 1.0
                    break
                elif pb in pi or pi in pb:
                    ratio_sub = min(len(pb), len(pi)) / max(len(pb), len(pi))
                    best_p = max(best_p, 0.88 * ratio_sub)
                else:
                    ratio_seq = difflib.SequenceMatcher(None, pb, pi).ratio()
                    best_p = max(best_p, ratio_seq)
            scores_palavras.append(best_p)
            
        score_cobertura = sum(scores_palavras) / len(scores_palavras) if scores_palavras else 0.0
        ratio_global = difflib.SequenceMatcher(None, termo_norm, item_norm).ratio()
        
        score_final = (score_cobertura * 0.40) + (score_ia * 0.35) + (score_substring * 0.15) + (ratio_global * 0.10)
        
        if all(any(pb == pi or pb in pi for pi in words_item) for pb in words_busca):
            score_final += 0.15
            
        if score_final > melhor_score_final:
            melhor_score_final = score_final
            melhor_item = item
            
    if melhor_score_final >= 0.35:
        return melhor_item, melhor_score_final
    return None, melhor_score_final

# ==============================================================================
# 3. BUSCA DOS DADOS NAVEGADOR / GOOGLE SHEETS
# ==============================================================================
@st.cache_data(ttl=300)
def buscar_itens_nuvem():
    try:
        with urllib.request.urlopen(URL_WEB_APP, timeout=10) as res:
            return json.loads(res.read().decode('utf-8'))
    except Exception as e:
        st.error(f"Erro ao carregar produtos da nuvem: {e}")
        return []

NOVOS_ITENS = buscar_itens_nuvem()
SETORES = ["RESTAURANTE / COZINHA", "BAR", "SALÃO"]

if 'usuario_anterior' not in st.session_state: st.session_state.usuario_anterior = ""
if 'carrinho_df' not in st.session_state: 
    st.session_state.carrinho_df = pd.DataFrame(columns=["Data_Hora", "Solicitante", "Setor", "Codigo", "Categoria", "Item", "Quantidade", "Unidade", "Observacao"])
if 'reset_counter' not in st.session_state: st.session_state.reset_counter = 0

# ==============================================================================
# INTERFACE PRINCIPAL
# ==============================================================================
st.title("📝 Sistema de Requisição")
nome_solicitante = st.text_input("Nome do Solicitante:")
setor_selecionado = st.selectbox("Selecione o seu Setor:", SETORES)

# Backup local por usuário
nome_limpo_idx = remover_acentos(nome_solicitante).replace(" ", "_") if nome_solicitante.strip() else ""

if nome_limpo_idx and st.session_state.usuario_anterior != nome_solicitante:
    arquivo_backup = f"backup_{nome_limpo_idx}.csv"
    if os.path.exists(arquivo_backup):
        try:
            df_backup = pd.read_csv(arquivo_backup)
            st.session_state.carrinho_df = df_backup
            st.toast(f"🔄 Pedido anterior de {nome_solicitante} foi recuperado!")
        except Exception:
            pass
    st.session_state.usuario_anterior = nome_solicitante

# Processamento da Lista Global do catálogo
mapeamento_itens = {}
for item in NOVOS_ITENS:
    desc = item.get("descricao", "").strip().upper()
    if desc:
        mapeamento_itens[desc] = {
            "codigo": str(item.get("codigo", "")).strip() if item.get("codigo") else "-",
            "categoria": str(item.get("categoria", "")).strip().upper() if item.get("categoria") else "OUTROS",
            "unidade": str(item.get("unidade", "")).strip().upper() if item.get("unidade") else "UND"
        }

opcoes_itens = sorted(list(mapeamento_itens.keys()))

# ==============================================================================
# 🎙️ MÓDULO DE FALA MULTI-ITENS INTELIGENTE COM IA
# ==============================================================================
st.write("### 🎙️ Ditar Múltiplos Itens por Voz (IA)")
st.caption("Exemplo de fala contínua: *'Quero 3 filé mignon, 2 heineken e 5 coca cola'*")

texto_falado = speech_to_text(
    language='pt-BR',
    start_prompt="🎙️ Clique para Ditar Vários Itens",
    stop_prompt="⏹️ Parar Gravação",
    just_once=True,
    key=f'stt_{st.session_state.reset_counter}'
)

if texto_falado:
    if not nome_solicitante.strip():
        st.error("⚠️ Preencha seu nome no campo acima antes de ditar os itens!")
    else:
        with st.spinner("🧠 Analisando fala e buscando nos produtos por completo..."):
            itens_extraidos = extrair_itens_da_fala(texto_falado)
            embeddings_cat = calcular_embeddings_catalogo(modelo_ia, opcoes_itens)
            
            adicionados = []
            nao_encontrados = []
            
            for orig, qtd_detectada, termo_busca in itens_extraidos:
                item_encontrado, score = encontrar_item_super_assertivo(termo_busca, opcoes_itens, embeddings_cat)
                
                if item_encontrado:
                    dados = mapeamento_itens[item_encontrado]
                    cod, cat, uni = dados["codigo"], dados["categoria"], dados["unidade"]
                    texto_obs = "-"
                    
                    mask = (st.session_state.carrinho_df["Item"] == item_encontrado) & \
                           (st.session_state.carrinho_df["Codigo"] == cod) & \
                           (st.session_state.carrinho_df["Observacao"] == texto_obs)
                           
                    if mask.any():
                        idx = st.session_state.carrinho_df[mask].index[0]
                        st.session_state.carrinho_df.at[idx, "Quantidade"] += qtd_detectada
                    else:
                        novo = pd.DataFrame([{
                            "Data_Hora": datetime.now().strftime("%d/%m/%Y %H:%M"), 
                            "Solicitante": nome_solicitante, 
                            "Setor": setor_selecionado, 
                            "Codigo": cod,
                            "Categoria": cat, 
                            "Item": item_encontrado, 
                            "Quantidade": qtd_detectada, 
                            "Unidade": uni, 
                            "Observacao": texto_obs
                        }])
                        st.session_state.carrinho_df = pd.concat([st.session_state.carrinho_df, novo], ignore_index=True)
                    
                    adicionados.append(f"{qtd_detectada}x {item_encontrado}")
                else:
                    nao_encontrados.append(termo_busca)
                    
            if adicionados:
                if nome_limpo_idx:
                    st.session_state.carrinho_df.to_csv(f"backup_{nome_limpo_idx}.csv", index=False)
                for item_msg in adicionados:
                    st.toast(f"✅ Adicionado: {item_msg}")
                    
            if nao_encontrados:
                st.warning(f"⚠️ Não foi possível identificar o(s) termo(s): {', '.join(nao_encontrados)}")
                
            if adicionados:
                st.session_state.reset_counter += 1
                time.sleep(1)
                st.rerun()

st.write("---")

# ==============================================================================
# 📄 MÓDULO DE ANEXAR DOCUMENTO (EXCEL, WORD, PDF)
# ==============================================================================
with st.expander("📄 **Anexar Arquivo com Lista de Itens (Excel, Word, PDF)**", expanded=False):
    st.caption("Faça o upload de uma planilha, documento do Word ou PDF com a lista de itens desejados.")
    
    arquivo_enviado = st.file_uploader(
        "Escolha o arquivo:", 
        type=["xlsx", "xls", "docx", "pdf"],
        key=f"uploader_{st.session_state.reset_counter}"
    )
    
    if arquivo_enviado is not None:
        if st.button("🔍 Processar Arquivo e Adicionar ao Carrinho", use_container_width=True):
            if not nome_solicitante.strip():
                st.error("⚠️ Preencha seu nome no campo acima antes de importar o arquivo!")
            else:
                with st.spinner("⚙️ Lendo documento e comparando com o catálogo..."):
                    linhas_documento = extrair_linhas_do_arquivo(arquivo_enviado)
                    embeddings_cat = calcular_embeddings_catalogo(modelo_ia, opcoes_itens)
                    
                    total_adicionados = 0
                    termos_nao_encontrados = []
                    
                    for linha in linhas_documento:
                        # Extrai os itens e quantidades presentes em cada linha do documento
                        itens_extraidos = extrair_itens_da_fala(linha)
                        
                        for orig, qtd_detectada, termo_busca in itens_extraidos:
                            item_encontrado, score = encontrar_item_super_assertivo(termo_busca, opcoes_itens, embeddings_cat)
                            
                            if item_encontrado:
                                dados = mapeamento_itens[item_encontrado]
                                cod, cat, uni = dados["codigo"], dados["categoria"], dados["unidade"]
                                texto_obs = f"Importado de {arquivo_enviado.name}"
                                
                                mask = (st.session_state.carrinho_df["Item"] == item_encontrado) & \
                                       (st.session_state.carrinho_df["Codigo"] == cod) & \
                                       (st.session_state.carrinho_df["Observacao"] == texto_obs)
                                       
                                if mask.any():
                                    idx = st.session_state.carrinho_df[mask].index[0]
                                    st.session_state.carrinho_df.at[idx, "Quantidade"] += qtd_detectada
                                else:
                                    novo = pd.DataFrame([{
                                        "Data_Hora": datetime.now().strftime("%d/%m/%Y %H:%M"), 
                                        "Solicitante": nome_solicitante, 
                                        "Setor": setor_selecionado, 
                                        "Codigo": cod,
                                        "Categoria": cat, 
                                        "Item": item_encontrado, 
                                        "Quantidade": qtd_detectada, 
                                        "Unidade": uni, 
                                        "Observacao": texto_obs
                                    }])
                                    st.session_state.carrinho_df = pd.concat([st.session_state.carrinho_df, novo], ignore_index=True)
                                
                                total_adicionados += 1
                            else:
                                if termo_busca not in termos_nao_encontrados:
                                    termos_nao_encontrados.append(termo_busca)
                    
                    if total_adicionados > 0:
                        if nome_limpo_idx:
                            st.session_state.carrinho_df.to_csv(f"backup_{nome_limpo_idx}.csv", index=False)
                        st.toast(f"✅ {total_adicionados} item(ns) importado(s) com sucesso!")
                        
                    if termos_nao_encontrados:
                        st.warning(f"⚠️ Não foi possível mapear o(s) seguinte(s) termo(s) do arquivo: {', '.join(termos_nao_encontrados)}")
                        
                    if total_adicionados > 0:
                        st.session_state.reset_counter += 1
                        time.sleep(1)
                        st.rerun()

st.write("---")

# ==============================================================================
# 4. SELEÇÃO MANUAL / FORMULÁRIO
# ==============================================================================
item_nao_cadastrado = st.checkbox("⚠️ Item NÃO está na lista?", key=f"manual_{st.session_state.reset_counter}")

if item_nao_cadastrado:
    item_bruto = st.text_input("Nome do Item:", key=f"item_manual_{st.session_state.reset_counter}")
    unidade_medida = st.selectbox("Unidade:", ["UND", "KG", "L", "LATA", "PORÇÃO", "G", "ML", "GFA", "PCT", "MAÇO", "CX"], key=f"uni_manual_{st.session_state.reset_counter}")
    codigo_detectado = st.text_input("Código (Se houver):", value="-", key=f"cod_manual_{st.session_state.reset_counter}")
    item_nome_limpo = item_bruto.upper().strip()
    subcategoria_detectada = "MANUAL"
else:
    if opcoes_itens:
        opcoes_display = [f"{i} ({remover_acentos(i)})" if remover_acentos(i) != i else i for i in opcoes_itens]
        mapeamento_display = {f"{i} ({remover_acentos(i)})" if remover_acentos(i) != i else i: i for i in opcoes_itens}
        
        item_bruto_display = st.selectbox(
            "Selecione o Item:", 
            opcoes_display, 
            index=None, 
            placeholder="Digite o nome do item para buscar...", 
            key=f"item_select_{st.session_state.reset_counter}"
        )
        
        if item_bruto_display:
            item_original_desc = mapeamento_display[item_bruto_display]
            dados_item = mapeamento_itens[item_original_desc]
            item_nome_limpo = item_original_desc
            unidade_medida = dados_item["unidade"]
            codigo_detectado = dados_item["codigo"]
            subcategoria_detectada = dados_item["categoria"]
            
            st.markdown(f"⚖️ **Unidade de Medida Padrão:** `{unidade_medida}`")
        else:
            item_nome_limpo = ""
            unidade_medida = "UND"
            codigo_detectado = "-"
            subcategoria_detectada = "OUTROS"
    else:
        st.warning("Carregando lista de produtos ou a planilha na nuvem está vazia...")
        item_nome_limpo = ""
        unidade_medida = "UND"
        codigo_detectado = "-"
        subcategoria_detectada = "OUTROS"

quantidade = st.number_input("Quantidade:", min_value=1, value=1, step=1, key=f"qtd_{st.session_state.reset_counter}")
observacao = st.text_input("Observação:", key=f"obs_{st.session_state.reset_counter}")

# ==============================================================================
# 5. LÓGICA DO CARRINHO E ENVIO PARA A CENTRAL
# ==============================================================================
if st.button("➕ Adicionar ao Pedido", use_container_width=True):
    if not nome_solicitante.strip(): st.error("Preencha seu nome.")
    elif not item_nome_limpo: st.error("Selecione ou digite um item válido.")
    else:
        texto_obs = observacao.strip() if observacao.strip() else "-"
        
        mask = (st.session_state.carrinho_df["Item"] == item_nome_limpo) & \
               (st.session_state.carrinho_df["Codigo"] == codigo_detectado) & \
               (st.session_state.carrinho_df["Observacao"] == texto_obs)
               
        if mask.any():
            idx = st.session_state.carrinho_df[mask].index[0]
            st.session_state.carrinho_df.at[idx, "Quantidade"] += int(quantidade)
        else:
            novo = pd.DataFrame([{
                "Data_Hora": datetime.now().strftime("%d/%m/%Y %H:%M"), 
                "Solicitante": nome_solicitante, 
                "Setor": setor_selecionado, 
                "Codigo": codigo_detectado,
                "Categoria": subcategoria_detectada, 
                "Item": item_nome_limpo, 
                "Quantidade": int(quantidade), 
                "Unidade": unidade_medida, 
                "Observacao": texto_obs
            }])
            st.session_state.carrinho_df = pd.concat([st.session_state.carrinho_df, novo], ignore_index=True)
        
        if nome_limpo_idx:
            st.session_state.carrinho_df.to_csv(f"backup_{nome_limpo_idx}.csv", index=False)
            
        st.session_state.reset_counter += 1
        st.rerun()

if not st.session_state.carrinho_df.empty:
    st.write("### 🛒 Pedido Atual")
    st.session_state.carrinho_df = st.data_editor(
        st.session_state.carrinho_df, 
        column_config={
            "Data_Hora": None, 
            "Solicitante": None, 
            "Setor": None, 
            "Codigo": st.column_config.TextColumn(disabled=True),
            "Categoria": None, 
            "Item": st.column_config.TextColumn(disabled=True), 
            "Unidade": st.column_config.TextColumn(disabled=True)
        }, 
        use_container_width=True, 
        num_rows="dynamic"
    )
    
    if nome_limpo_idx:
        st.session_state.carrinho_df.to_csv(f"backup_{nome_limpo_idx}.csv", index=False)
    
    if st.button("🚀 ENVIAR PARA A CENTRAL", type="primary", use_container_width=True):
        try:
            req = urllib.request.Request(
                URL_WEB_APP, 
                method="POST", 
                data=json.dumps(st.session_state.carrinho_df.to_dict(orient='records')).encode('utf-8'), 
                headers={'Content-Type': 'application/json'}
            )
            with urllib.request.urlopen(req) as res:
                if "Success" in res.read().decode('utf-8'):
                    st.balloons()
                    st.success("Enviado com sucesso!")
                    
                    if nome_limpo_idx and os.path.exists(f"backup_{nome_limpo_idx}.csv"):
                        os.remove(f"backup_{nome_limpo_idx}.csv")
                        
                    st.session_state.carrinho_df = pd.DataFrame(
                        columns=["Data_Hora", "Solicitante", "Setor", "Codigo", "Categoria", "Item", "Quantidade", "Unidade", "Observacao"]
                    )
                    time.sleep(2); st.rerun()
        except Exception as e: st.error(f"Erro: {e}")
